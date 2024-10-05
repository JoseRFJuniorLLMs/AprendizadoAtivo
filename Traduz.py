import docx
from collections import Counter
from googletrans import Translator
import nltk
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
import re

# Baixar recursos necessários do NLTK
nltk.download('punkt', quiet=True)
nltk.download('stopwords', quiet=True)

def read_docx(file_path):
    """Lê o arquivo .docx e retorna o texto completo."""
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return ' '.join(full_text)

def is_valid_word(word):
    """Verifica se a palavra é válida (somente alfanumérica e não contém caracteres especiais)."""
    return bool(re.match(r'^[A-Za-zÀ-ÿ0-9]+$', word))

def get_top_words(text, n=100):
    """Obtém as n palavras mais comuns do texto, excluindo stopwords."""
    tokens = word_tokenize(text.lower())
    stop_words = set(stopwords.words('portuguese'))
    tokens = [word for word in tokens if is_valid_word(word) and word not in stop_words]
    word_counts = Counter(tokens)
    return word_counts.most_common(n)

def translate_word(word):
    """Traduz uma única palavra usando o Google Translate."""
    translator = Translator()
    try:
        return translator.translate(word, src='pt', dest='en').text
    except Exception as e:
        print(f"Erro ao traduzir '{word}': {e}")
        return None

def replace_word_in_doc(doc, word, translation):
    """
    Substitui a palavra no documento pelo seu equivalente em inglês,
    garantindo que seja uma correspondência exata e ignorando maiúsculas/minúsculas.
    """
    for para in doc.paragraphs:
        for run in para.runs:
            if word.lower() in run.text.lower():
                # Realiza a substituição da palavra
                run.text = re.sub(rf'\b{word}\b', translation, run.text, flags=re.IGNORECASE)
    return doc

def process_file(file_path, num_words=200):
    """Processa o arquivo .docx especificado."""
    # Ler o documento
    text = read_docx(file_path)

    # Obter as n palavras mais comuns
    top_words = get_top_words(text, num_words)

    # Traduzir e substituir as palavras no documento
    new_doc = docx.Document(file_path)
    translated_words = {}
    for word, _ in top_words:
        if is_valid_word(word):
            translation = translate_word(word)
            if translation:
                translated_words[word] = translation
                new_doc = replace_word_in_doc(new_doc, word, translation)

    # Criar uma lista das palavras traduzidas
    translated_words_list = []
    for word, translation in translated_words.items():
        translated_words_list.append(f"{word} - {translation}")

    # Adicionar a lista de palavras traduzidas ao documento
    new_doc.add_paragraph("Lista de palavras traduzidas:")
    for translated_word in translated_words_list:
        new_doc.add_paragraph(translated_word)

    # Salvar o novo documento com o sufixo '-A0' na mesma pasta
    new_file_name = 'docsx/DomQuixote-A0.docx'
    new_doc.save(new_file_name)
    print(f"Processado: {file_path} -> {new_file_name}")

# Caminho do arquivo .docx a ser processado
file_path = 'DomQuixote.docx'  # Nome do arquivo fixo
num_words = 200  # Número de palavras mais comuns a serem obtidas

process_file(file_path, num_words)
print("Processo concluído.")
